# -*- coding: utf-8 -*-
"""Kit de geracao dos documentos da Explosao Junina de Beruri.

Padrao visual unico (Arial, vermelho 922B21, grafite 2C3E50) usado por todos os
`gen_*.py` desta pasta. Rodar cada gerador reescreve o .docx correspondente na
pasta "documentos explosao" — editar sempre o gerador, nunca o .docx na mao.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- paleta ----
VERMELHO = "922B21"   # marca / titulos
GRAFITE = "2C3E50"    # texto
AMBAR = "F39C12"      # destaque
LARANJA = "E67E22"
AZUL = "2980B9"
VERDE = "27AE60"
ROXO = "8E44AD"
CINZA_FUNDO = "F2F3F4"
CINZA_TEXTO = "7F8C8D"
BRANCO = "FFFFFF"

FONTE = "Arial"

CABECALHO = [
    ("QUADRILHA JUNINA", 22, True, VERMELHO),
    ("EXPLOSÃO JUNINA DE BERURI", 22, True, VERMELHO),
    ("Grupo Recreativo e Folclórico | Campeã do Festival Folclórico 2025", 10, False, GRAFITE),
    ("Filiada à LIQUAJUAM | Fundada em 02/06/2017", 9, False, GRAFITE),
]

PASTA_DOCS = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def rgb(hexcor):
    return RGBColor(int(hexcor[0:2], 16), int(hexcor[2:4], 16), int(hexcor[4:6], 16))


# ------------------------------------------------------------- documento ----
def novo_documento():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Cm(2.5))
    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(GRAFITE)
    return doc


def _run(par, texto, tam=11, bold=False, cor=GRAFITE, italico=False):
    r = par.add_run(texto)
    r.font.name = FONTE
    r.font.size = Pt(tam)
    r.font.bold = bold
    r.font.italic = italico
    r.font.color.rgb = rgb(cor)
    return r


def espaco(doc, pt=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    return p


def capa(doc, titulo, subtitulo, linha_apoio=None, nota=None, rodape=None, cor=GRAFITE):
    """Capa padrao: cabecalho da quadrilha + titulo grande + apoio."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(100)
    for texto, tam, bold, cor_txt in CABECALHO:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after = Pt(3 if tam < 20 else 3)
        _run(par, texto, tam, bold, cor_txt)
    espaco(doc, 20)
    for linha in (titulo if isinstance(titulo, (list, tuple)) else [titulo]):
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after = Pt(4)
        _run(par, linha, 26, True, cor)
    espaco(doc, 10)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(3)
    _run(par, subtitulo, 14, True, AMBAR)
    if linha_apoio:
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.space_after = Pt(3)
        _run(par, linha_apoio, 11, False, GRAFITE)
    if nota:
        espaco(doc, 30)
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(par, nota, 10, False, CINZA_TEXTO, italico=True)
    if rodape:
        espaco(doc, 40)
        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(par, rodape, 11, False, GRAFITE)
    doc.add_page_break()


def h1(doc, texto, numero=None):
    par = doc.add_paragraph(style="Heading 1")
    par.paragraph_format.space_before = Pt(18)
    par.paragraph_format.space_after = Pt(10)
    _run(par, (f"{numero}. {texto}" if numero else texto), 16, True, VERMELHO)
    return par


def h2(doc, texto):
    par = doc.add_paragraph(style="Heading 2")
    par.paragraph_format.space_before = Pt(14)
    par.paragraph_format.space_after = Pt(8)
    _run(par, texto, 13, True, GRAFITE)
    return par


def h3(doc, texto):
    par = doc.add_paragraph(style="Heading 3")
    par.paragraph_format.space_before = Pt(10)
    par.paragraph_format.space_after = Pt(4)
    _run(par, texto, 11.5, True, VERMELHO)
    return par


def p(doc, texto, bold_ate=None, tam=11, cor=GRAFITE, italico=False, centro=False):
    """Paragrafo justificado. `bold_ate` = trecho inicial em negrito (ex.: rotulo)."""
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if centro else WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.space_after = Pt(7)
    if bold_ate:
        _run(par, bold_ate, tam, True, cor)
    _run(par, texto, tam, False, cor, italico)
    return par


def bullet(doc, texto, rotulo=None, tam=11):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(4)
    if rotulo:
        _run(par, rotulo, tam, True, GRAFITE)
    _run(par, texto, tam, False, GRAFITE)
    return par


def numero(doc, texto, rotulo=None, tam=11):
    par = doc.add_paragraph(style="List Number")
    par.paragraph_format.space_after = Pt(4)
    if rotulo:
        _run(par, rotulo, tam, True, GRAFITE)
    _run(par, texto, tam, False, GRAFITE)
    return par


def citacao(doc, texto):
    espaco(doc, 6)
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(10)
    par.paragraph_format.space_after = Pt(10)
    _run(par, f"“{texto}”", 12, True, VERMELHO, italico=True)
    return par


# ---------------------------------------------------------------- shading ---
def _fundo(elemento, cor_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor_hex)
    elemento.append(shd)


def _borda_tabela(tabela, cor=CINZA_TEXTO, sz=4):
    tbl_pr = tabela._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for lado in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{lado}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), cor)
        borders.append(el)
    tbl_pr.append(borders)


def _celula(cell, texto, tam=10, bold=False, cor=GRAFITE, fundo=None, centro=False):
    if fundo:
        _fundo(cell._tc.get_or_add_tcPr(), fundo)
    cell.text = ""
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if centro else WD_ALIGN_PARAGRAPH.LEFT
    par.paragraph_format.space_before = Pt(3)
    par.paragraph_format.space_after = Pt(3)
    linhas = texto.split("\n") if isinstance(texto, str) else [str(texto)]
    for i, linha in enumerate(linhas):
        if i:
            par = cell.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER if centro else WD_ALIGN_PARAGRAPH.LEFT
            par.paragraph_format.space_before = Pt(0)
            par.paragraph_format.space_after = Pt(3)
        _run(par, linha, tam, bold, cor)


def banner(doc, titulo, subtitulo="", cor=VERMELHO):
    """Faixa colorida de abertura de secao/ficha."""
    espaco(doc, 6)
    tb = doc.add_table(rows=1, cols=1)
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tb.cell(0, 0)
    _fundo(cell._tc.get_or_add_tcPr(), cor)
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(2 if subtitulo else 6)
    _run(par, titulo, 15, True, BRANCO)
    if subtitulo:
        par2 = cell.add_paragraph()
        par2.paragraph_format.space_after = Pt(6)
        _run(par2, subtitulo, 10, False, BRANCO)
    espaco(doc, 6)
    return tb


def caixa(doc, titulo, texto, cor=AMBAR):
    """Caixa de destaque (fundo cinza, titulo colorido)."""
    tb = doc.add_table(rows=1, cols=1)
    cell = tb.cell(0, 0)
    _fundo(cell._tc.get_or_add_tcPr(), CINZA_FUNDO)
    cell.text = ""
    par = cell.paragraphs[0]
    par.paragraph_format.space_before = Pt(5)
    par.paragraph_format.space_after = Pt(2)
    _run(par, titulo, 11, True, cor)
    par2 = cell.add_paragraph()
    par2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par2.paragraph_format.space_after = Pt(6)
    _run(par2, texto, 10.5, False, GRAFITE)
    _borda_tabela(tb, cor)
    espaco(doc, 6)
    return tb


def tabela(doc, cabecalhos, linhas, larguras=None, cor_cabecalho=VERMELHO, tam=10):
    tb = doc.add_table(rows=1, cols=len(cabecalhos))
    tb.alignment = WD_TABLE_ALIGNMENT.CENTER
    tb.autofit = False
    for i, texto in enumerate(cabecalhos):
        _celula(tb.rows[0].cells[i], texto, tam, True, BRANCO, cor_cabecalho)
    for linha in linhas:
        cells = tb.add_row().cells
        for i, valor in enumerate(linha):
            bold = (i == 0 and len(linha) > 1)
            _celula(cells[i], valor, tam, bold, GRAFITE, CINZA_FUNDO)
    if larguras:
        for row in tb.rows:
            for i, larg in enumerate(larguras):
                row.cells[i].width = Cm(larg)
    _borda_tabela(tb)
    espaco(doc, 8)
    return tb


def checklist(doc, itens, titulo="Checklist de execução"):
    if titulo:
        h2(doc, titulo)
    tabela(doc, ["O que fazer", "Responsável", "Status"],
           [[i, "", "[    ]"] for i in itens],
           larguras=[10.5, 3.5, 2.0])


def salvar(doc, nome_arquivo):
    caminho = os.path.join(PASTA_DOCS, nome_arquivo)
    doc.save(caminho)
    print("OK ->", caminho)
    return caminho
