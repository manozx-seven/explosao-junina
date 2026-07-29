# -*- coding: utf-8 -*-
"""Insere o TETO DE R$ 80,00 da bonificacao no 'Contratos Explosao Junina Final.docx'.

Por que este script existe (e nao um gerador completo): o contrato e o unico
documento da pasta que ainda nao tem `gen_*.py`, e reescreve-lo do zero para
mudar uma regra arriscaria perder formatacao juridica (campos de assinatura,
caixas de opcao, anexos). Entao a alteracao entra cirurgicamente — mas por
script, versionado, e nunca na mao: continua valendo a regra do projeto de que
nenhum .docx e editado direto.

O que faz, nos DOIS termos (Brincante e Item Destaque):
  1. Clausula Sexta, III — acrescenta as alineas "e" (o teto) e "f" (por que ele
     existe e o que ele nao muda), logo depois da alinea "d".
  2. Clausula Sexta, X — acrescenta a alinea que amarra o Programa de Fidelidade
     ao mesmo teto (Nivel 3 dobra o valor por ensaio; sem esta linha, um
     documento diria uma coisa e o outro, outra).
  3. Projecao financeira — linha do teto na tabela e uma observacao explicando
     que, com os valores atuais, o maximo (R$ 44,00) ainda fica abaixo dele.
  4. Quadro de adesao da primeira pagina — o teto aparece onde o brincante
     assina, nao so no meio do contrato.

E idempotente: rodar duas vezes nao duplica nada.

Uso:  python patch_contrato_teto.py
"""

import copy
import os
import sys

from docx import Document

PASTA_DOCS = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
ARQUIVO = os.path.join(PASTA_DOCS, "Contratos Explosao Junina Final.docx")

TETO = "R$ 80,00"

ALINEA_E = (
    "e) ",
    "TETO DA BONIFICAÇÃO: o valor acumulado por brincante é limitado a "
    f"{TETO} (oitenta reais) por temporada. Ao atingir o teto, o brincante "
    "permanece com o direito ao valor acumulado até ali, mas não acumula "
    "novos valores até o fim da temporada;",
)

ALINEA_F = (
    "f) ",
    "O teto existe para que o Programa cumpra seu papel de retorno simbólico "
    "sem extrapolar o orçamento da quadrilha. É o mesmo para todos, é "
    "acompanhado em tempo real no sistema de avaliação — que mostra quanto "
    "falta para alcançá-lo — e não altera nenhuma das demais regras desta "
    "Cláusula: frequência, desempenho, ativação, sanções e condições de "
    "resgate continuam valendo integralmente. Atingir o teto não dispensa o "
    "brincante de nada;",
)

ALINEA_FIDELIDADE = (
    "f) ",
    f"Em qualquer nível do Programa de Fidelidade, o valor acumulado continua "
    f"limitado ao teto de {TETO} por temporada, previsto na seção III, alínea "
    "'e' desta Cláusula.",
)

OBS_TETO = (
    "Obs. 2: ",
    f"Nenhum brincante acumula mais de {TETO} na temporada (teto da seção III, "
    "alínea 'e'). Com os valores atuais, o máximo possível é R$ 44,00 — o teto "
    "só passa a limitar de fato se os valores por evento forem revistos para "
    "cima ou se a temporada tiver muito mais apresentações do que o previsto. "
    "Ele é a garantia de que a bonificação nunca estoura o orçamento, "
    "qualquer que seja o tamanho da temporada.",
)

FRASE_ADESAO_DE = (
    "Ao optar por participar, o brincante aceita cumprir os critérios de "
    "ativação e desempenho descritos na Cláusula Sexta."
)
FRASE_ADESAO_PARA = (
    "Ao optar por participar, o brincante aceita cumprir os critérios de "
    "ativação e desempenho descritos na Cláusula Sexta. O valor acumulado tem "
    f"teto de {TETO} (oitenta reais) por temporada."
)


def inserir_apos(par_modelo, rotulo, texto):
    """Clona o paragrafo-modelo e o insere logo abaixo, so trocando o texto.

    Clonar em vez de criar do zero e o que mantem fonte, tamanho, recuo e
    espacamento identicos aos das outras alineas — o contrato e um documento
    que sera impresso e assinado, e uma alinea com cara diferente denuncia
    remendo.
    """
    novo = copy.deepcopy(par_modelo._p)
    par_modelo._p.addnext(novo)

    from docx.text.paragraph import Paragraph
    par = Paragraph(novo, par_modelo._parent)

    # O modelo tem dois runs: o rotulo em negrito e o corpo em texto normal.
    runs = par.runs
    if len(runs) >= 2:
        runs[0].text = rotulo
        runs[1].text = texto
        for r in runs[2:]:
            r._r.getparent().remove(r._r)
    else:
        runs[0].text = rotulo + texto
    return par


def clonar_ultima_linha(tabela, valores):
    """Acrescenta uma linha copiando a ultima (mesma formatacao de celula)."""
    nova = copy.deepcopy(tabela.rows[-1]._tr)
    tabela.rows[-1]._tr.addnext(nova)
    linha = tabela.rows[-1]
    for cel, valor in zip(linha.cells, valores):
        par = cel.paragraphs[0]
        for extra in cel.paragraphs[1:]:
            extra._p.getparent().remove(extra._p)
        if par.runs:
            par.runs[0].text = valor
            for r in par.runs[1:]:
                r._r.getparent().remove(r._r)
        else:
            par.add_run(valor)
    return linha


def main():
    if not os.path.exists(ARQUIVO):
        print("nao achei:", ARQUIVO)
        return 1

    doc = Document(ARQUIVO)

    if any(TETO in p.text for p in doc.paragraphs):
        print("O teto ja esta no documento — nada a fazer.")
        return 0

    feitos = {"alineas": 0, "fidelidade": 0, "obs": 0, "tabela": 0, "adesao": 0}

    # 1 e 2 — alineas novas. Percorre de tras para frente porque inserir
    # paragrafos desloca os indices seguintes.
    for par in reversed(doc.paragraphs):
        texto = par.text.strip()
        if texto.startswith("d) Os valores serão registrados no sistema"):
            inserir_apos(par, *ALINEA_F)
            inserir_apos(par, *ALINEA_E)
            feitos["alineas"] += 1
        elif texto.startswith("e) Os valores e benefícios por nível poderão ser ajustados"):
            inserir_apos(par, *ALINEA_FIDELIDADE)
            feitos["fidelidade"] += 1
        elif texto.startswith("Obs.: Os valores de apresentações externas"):
            inserir_apos(par, *OBS_TETO)
            feitos["obs"] += 1

    # 3 — linha do teto na tabela de projecao por brincante.
    for tb in doc.tables:
        primeira = " ".join(c.text for c in tb.rows[0].cells)
        if "Unitário" in primeira and "Total/brincante" in primeira:
            clonar_ultima_linha(tb, [
                "TETO DA TEMPORADA (limite máximo acumulável)", "—", "—", TETO,
            ])
            feitos["tabela"] += 1

    # 4 — quadro de adesao da primeira pagina.
    for tb in doc.tables:
        for linha in tb.rows:
            for cel in linha.cells:
                for par in cel.paragraphs:
                    if FRASE_ADESAO_DE not in par.text:
                        continue
                    alvo = next((r for r in par.runs if FRASE_ADESAO_DE in r.text), None)
                    if alvo is not None:
                        alvo.text = alvo.text.replace(FRASE_ADESAO_DE, FRASE_ADESAO_PARA)
                    else:
                        # O Word quebrou a frase em varios runs: reescreve o
                        # paragrafo inteiro no primeiro run, que carrega a
                        # formatacao, e descarta os demais.
                        novo_texto = par.text.replace(FRASE_ADESAO_DE, FRASE_ADESAO_PARA)
                        par.runs[0].text = novo_texto
                        for r in par.runs[1:]:
                            r._r.getparent().remove(r._r)
                    feitos["adesao"] += 1

    doc.save(ARQUIVO)
    print("OK ->", ARQUIVO)
    for chave, n in feitos.items():
        print(f"   {chave}: {n}x")
    if feitos["alineas"] != 2:
        print("   AVISO: esperava 2 ocorrencias (Brincante e Item Destaque).")
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
